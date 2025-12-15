"""
XMind -> DOCX 转换脚本
保留目录（Word 导航/TOC）、文字和图片。
"""

import argparse
import imghdr
import json
import sys
import typing as typing
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Dict, Tuple

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches
except ImportError:
    print("缺少依赖 python-docx，请先安装：pip install python-docx")
    sys.exit(1)

try:
    import xmind  # 作为旧格式回退方案
except ImportError:
    xmind = None


def add_toc(document: Document):
    """插入 TOC 域，需在 Word 中更新域后生效。"""
    p = document.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    p._p.append(fld)


def extract_resources(zf: zipfile.ZipFile) -> Dict[str, bytes]:
    """提取 resources 下的文件为内存字典，键为 src 路径。"""
    mapping: Dict[str, bytes] = {}
    for name in zf.namelist():
        if not name.startswith("resources/"):
            continue
        if name.endswith("/"):
            continue
        data = zf.read(name)
        mapping[name] = data
        mapping["xap:" + name] = data  # 兼容前缀
    return mapping


def load_from_content_json(path: Path) -> Tuple[typing.Optional[Dict], Dict[str, bytes]]:
    """优先解析新版 XMind (content.json)。"""
    try:
        with zipfile.ZipFile(path, "r") as zf:
            names = zf.namelist()
            candidates = [n for n in names if n.endswith("content.json")]
            if not candidates:
                return None, {}
            assets = extract_resources(zf)
            data = zf.read(candidates[0])
            doc = json.loads(data)
            sheet = doc[0]
            return sheet.get("rootTopic"), assets
    except Exception:
        return None, {}


def load_from_xmind_lib(path: Path) -> Tuple[typing.Optional[Dict], Dict[str, bytes]]:
    """回退使用 xmind 库解析老格式。"""
    if xmind is None:
        return None, {}
    workbook = xmind.load(str(path))
    sheet = workbook.getPrimarySheet()
    dictSheet = sheet.getData()
    return dictSheet.get("topic"), {}


def iter_children(topic: Dict):
    """统一遍历子节点，兼容老格式和新格式。"""
    for group in topic.get("topics", {}).values():
        for child in group:
            yield child
    children = topic.get("children", {})
    for key in ("attached", "detached"):
        for child in children.get(key, []):
            yield child


def add_topic(
    document: Document,
    topic: Dict,
    assets: Dict[str, bytes],
    level: int = 1,
    img_width_inch: float = 6.0,
):
    """递归写入 topic -> DOCX。level 从 1 起，对应 Heading 1。叶子节点写正文列表。"""
    raw_title = str(topic.get("title", "") or "")
    title = "".join(ch for ch in raw_title if (ch >= " " or ch in "\t\n\r")).strip() or "."
    children = list(iter_children(topic))
    is_leaf = len(children) == 0

    if level == 1 or not is_leaf:
        heading_level = min(level, 9)
        document.add_heading(title, level=heading_level)
    else:
        p = document.add_paragraph(title, style="List Bullet")
        try:
            p.paragraph_format.left_indent = Inches(0.25 * max(level - 1, 0))
        except Exception:
            pass  # 若样式不可用则忽略缩进

    # 插入图片（如果有）
    image = topic.get("image") or {}
    src = image.get("src")
    if src:
        data = assets.get(src)
        if data:
            try:
                kind = imghdr.what(None, h=data)
                if kind in {"png", "jpeg", "bmp", "gif"}:
                    document.add_picture(BytesIO(data), width=Inches(img_width_inch))
                else:
                    # 忽略不被 python-docx 识别的资源
                    pass
            except Exception:
                # 避免单个资源导致整体失败
                pass

    for child in children:
        add_topic(document, child, assets, level=level + 1, img_width_inch=img_width_inch)


def parse_args():
    parser = argparse.ArgumentParser(description="Convert XMind to DOCX (keep text, images, TOC).")
    parser.add_argument("source", help="XMind 文件路径")
    parser.add_argument("-o", "--output", help="输出 DOCX 路径（默认为同名 .docx）")
    parser.add_argument(
        "--img-width",
        type=float,
        default=6.0,
        help="插入图片的宽度（英寸），默认 6.0",
    )
    parser.add_argument("--no-toc", action="store_true", help="不插入 TOC 域")
    return parser.parse_args()


def main():
    args = parse_args()

    pathSource = args.source.strip('"').strip("'")
    source_path = Path(pathSource.replace("\\", "/")).expanduser()
    if not source_path.exists():
        print(f"未找到文件: {source_path}")
        sys.exit(1)

    if args.output:
        output_path = Path(args.output).expanduser()
    else:
        output_path = source_path.with_suffix(".docx")

    if output_path.parent == Path(""):
        output_path = Path.cwd() / output_path

    # 解析 XMind
    root_topic, assets = load_from_content_json(source_path)
    if root_topic is None:
        root_topic, assets = load_from_xmind_lib(source_path)
    if root_topic is None:
        print("无法解析 XMind 文件：缺少 content.json 且 xmind 库不可用。")
        sys.exit(1)

    # 构建 DOCX
    document = Document()
    document.core_properties.title = str(root_topic.get("title", "XMind"))
    if not args.no_toc:
        add_toc(document)
    add_topic(document, root_topic, assets, level=1, img_width_inch=max(args.img_width, 0.1))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(f"已生成 DOCX: {output_path}")


if __name__ == "__main__":
    main()
